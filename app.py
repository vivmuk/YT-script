import os
import re
import io
from flask import Flask, request, jsonify, render_template, send_file
import yt_dlp
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/transcribe', methods=['POST'])
def transcribe_video():
    try:
        data = request.get_json()
        youtube_url = data.get('url', '').strip()
        
        if not youtube_url:
            return jsonify({'error': 'No URL provided'}), 400
        
        # Validate YouTube URL
        if not is_valid_youtube_url(youtube_url):
            return jsonify({'error': 'Invalid YouTube URL'}), 400
        
        logger.info(f"Processing YouTube URL: {youtube_url}")
        
        # Extract closed captions from YouTube video
        captions_data = extract_youtube_captions(youtube_url)
        
        if not captions_data:
            return jsonify({'error': 'No closed captions found for this video. The video may not have captions available.'}), 404
        
        # Extract video info for context
        video_info = get_video_info(youtube_url)
        
        response_data = {
            'transcription': captions_data['text'],
            'title': video_info.get('title', 'Unknown Title'),
            'duration': video_info.get('duration', 'Unknown Duration'),
            'uploader': video_info.get('uploader', 'Unknown'),
            'language': captions_data.get('language', 'Unknown'),
            'caption_type': captions_data.get('type', 'Unknown'),
            'success': True
        }
        
        logger.info("Captions extracted successfully")
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error during caption extraction: {str(e)}")
        return jsonify({'error': f'Caption extraction failed: {str(e)}'}), 500

@app.route('/download/json', methods=['POST'])
def download_json():
    try:
        data = request.get_json()
        transcription = data.get('transcription', '')
        meta = {
            'title': data.get('title', ''),
            'duration': data.get('duration', ''),
            'uploader': data.get('uploader', ''),
            'language': data.get('language', ''),
            'caption_type': data.get('caption_type', ''),
        }

        payload = {
            'transcription': transcription,
            'metadata': meta,
        }

        import json
        json_bytes = json.dumps(payload, ensure_ascii=False, indent=2).encode('utf-8')
        filename = f"{slugify(meta.get('title') or 'youtube_captions')}.json"
        return send_file(
            io.BytesIO(json_bytes),
            mimetype='application/json',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"Error generating JSON: {str(e)}")
        return jsonify({'error': 'Failed to generate JSON'}), 500

@app.route('/download/docx', methods=['POST'])
def download_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches

        data = request.get_json()
        transcription = data.get('transcription', '')
        title = data.get('title', 'YouTube Captions')
        duration = data.get('duration', '')
        uploader = data.get('uploader', '')
        language = data.get('language', '')
        caption_type = data.get('caption_type', '')

        document = Document()
        heading = document.add_heading(title, level=0)
        heading.style.font.size = Pt(18)

        meta_para = document.add_paragraph()
        meta_para.add_run(f"Uploader: {uploader}\n").bold = True
        meta_para.add_run(f"Duration: {duration}\n").bold = True
        meta_para.add_run(f"Language: {language} | Type: {caption_type}\n").bold = True

        document.add_paragraph("")
        document.add_heading('Captions', level=1)

        # Write transcription in paragraphs, splitting on sentence-ish breaks
        for chunk in split_into_paragraphs(transcription):
            document.add_paragraph(chunk)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"{slugify(title or 'youtube_captions')}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        return jsonify({'error': 'Failed to generate DOCX'}), 500

def split_into_paragraphs(text: str):
    if not text:
        return []
    # Split on double newlines or sentence boundaries for readability
    parts = re.split(r"\n\n+|(?<=[.!?])\s+(?=[A-Z])", text)
    return [p.strip() for p in parts if p.strip()]

def slugify(value: str) -> str:
    value = value or ''
    value = value.lower()
    value = re.sub(r"[^a-z0-9\-\_\s]", '', value)
    value = re.sub(r"\s+", '_', value).strip('_')
    return value or 'file'

def is_valid_youtube_url(url):
    """Check if the URL is a valid YouTube URL"""
    youtube_domains = ['youtube.com', 'youtu.be', 'www.youtube.com', 'm.youtube.com']
    return any(domain in url for domain in youtube_domains)

def extract_youtube_captions(url):
    """Extract closed captions from YouTube video using yt-dlp"""
    try:
        ydl_opts = {
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitleslangs': ['en', 'en-US', 'en-GB'],  # Prefer English
            'skip_download': True,
            'quiet': True,
            'no_warnings': True,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            # Extract video info and subtitles
            info = ydl.extract_info(url, download=False)
            
            # Get available subtitles
            subtitles = info.get('subtitles', {})
            automatic_captions = info.get('automatic_captions', {})
            
            # Try to get the best available captions
            caption_data = None
            caption_type = None
            language = None
            
            # Priority: manual subtitles first, then automatic
            for lang_code in ['en', 'en-US', 'en-GB']:
                if lang_code in subtitles:
                    caption_data = subtitles[lang_code]
                    caption_type = 'Manual'
                    language = lang_code
                    break
            
            if not caption_data:
                for lang_code in ['en', 'en-US', 'en-GB']:
                    if lang_code in automatic_captions:
                        caption_data = automatic_captions[lang_code]
                        caption_type = 'Automatic'
                        language = lang_code
                        break
            
            # If no English captions, try any available language
            if not caption_data:
                if subtitles:
                    first_lang = list(subtitles.keys())[0]
                    caption_data = subtitles[first_lang]
                    caption_type = 'Manual'
                    language = first_lang
                elif automatic_captions:
                    first_lang = list(automatic_captions.keys())[0]
                    caption_data = automatic_captions[first_lang]
                    caption_type = 'Automatic'
                    language = first_lang
            
            if not caption_data:
                return None
            
            # Find the best format (prefer json3 or vtt, then others)
            best_format = None
            preferred_exts = ['json3', 'vtt', 'ttml', 'srv3', 'srv2', 'srv1']
            for fmt in caption_data:
                if fmt.get('ext') in preferred_exts:
                    best_format = fmt
                    break
            
            if not best_format:
                best_format = caption_data[0]  # Fallback to first available
            
            # Download the caption content using yt-dlp urlopen to include headers/cookies
            caption_url = best_format['url']
            caption_bytes = ydl.urlopen(caption_url).read()
            caption_content = caption_bytes.decode('utf-8', errors='ignore')
            
            # Parse the caption content
            parsed_text = parse_caption_content(caption_content, best_format.get('ext', ''))
            
            return {
                'text': parsed_text,
                'language': language,
                'type': caption_type,
                'format': best_format['ext']
            }
            
    except Exception as e:
        logger.error(f"Error extracting captions: {str(e)}")
        return None

def parse_caption_content(content, format_type):
    """Parse caption content based on format type"""
    try:
        if format_type == 'vtt':
            return parse_vtt_content(content)
        elif format_type == 'json3':
            return parse_json3_content(content)
        elif format_type in ['ttml', 'srv3', 'srv2', 'srv1']:
            return parse_xml_content(content)
        else:
            # Fallback: try to extract text from any format
            return clean_caption_text(content)
    except Exception as e:
        logger.error(f"Error parsing caption content: {str(e)}")
        return clean_caption_text(content)

def parse_vtt_content(vtt_content):
    """Parse WebVTT format captions"""
    lines = vtt_content.split('\n')
    text_lines = []
    
    for line in lines:
        line = line.strip()
        # Skip empty lines, timestamps, and WebVTT headers
        if (line and 
            not line.startswith('WEBVTT') and 
            not line.startswith('NOTE') and 
            not '-->' in line and 
            not line.isdigit() and
            not line.startswith('STYLE') and
            not line.startswith('::cue')):
            
            # Clean HTML tags and formatting
            clean_line = re.sub(r'<[^>]+>', '', line)
            clean_line = re.sub(r'&[a-zA-Z]+;', ' ', clean_line)  # Remove HTML entities
            clean_line = clean_line.strip()
            
            if clean_line:
                text_lines.append(clean_line)
    
    return ' '.join(text_lines)

def parse_json3_content(json_content):
    """Parse YouTube json3 automatic caption format"""
    try:
        import json
        # Some responses are prefixed with )]}' to prevent XSSI
        cleaned = json_content.lstrip() \
            .removeprefix(")]}'") \
            .lstrip()  # remove any leading newlines/spaces
        data = json.loads(cleaned)
        texts = []
        for event in data.get('events', []):
            for seg in event.get('segs', []) or []:
                t = seg.get('utf8', '')
                if t:
                    texts.append(t)
        # Join and normalize whitespace/newlines
        text = ''.join(texts)
        text = text.replace('\n', ' ')
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    except Exception:
        return clean_caption_text(json_content)

def parse_xml_content(xml_content):
    """Parse XML-based caption formats (TTML, SRV, etc.)"""
    # Remove XML tags and extract text content
    text = re.sub(r'<[^>]+>', ' ', xml_content)
    # Clean up HTML entities
    text = re.sub(r'&[a-zA-Z]+;', ' ', text)
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def clean_caption_text(content):
    """Fallback method to clean any caption format"""
    # Remove common caption formatting
    text = re.sub(r'<[^>]+>', ' ', content)  # Remove HTML/XML tags
    text = re.sub(r'&[a-zA-Z]+;', ' ', text)  # Remove HTML entities
    text = re.sub(r'\d{2}:\d{2}:\d{2}[.,]\d{3}', ' ', text)  # Remove timestamps
    text = re.sub(r'-->', ' ', text)  # Remove arrow separators
    text = re.sub(r'WEBVTT', ' ', text)  # Remove WebVTT headers
    text = re.sub(r'NOTE.*', ' ', text)  # Remove NOTE lines
    text = re.sub(r'\s+', ' ', text).strip()  # Clean up whitespace
    return text

def get_video_info(url):
    """Get video information without downloading"""
    try:
        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            return {
                'title': info.get('title', 'Unknown Title'),
                'duration': format_duration(info.get('duration', 0)),
                'uploader': info.get('uploader', 'Unknown'),
            }
    except Exception as e:
        logger.error(f"Error getting video info: {str(e)}")
        return {'title': 'Unknown Title', 'duration': 'Unknown Duration'}

def format_duration(seconds):
    """Format duration from seconds to HH:MM:SS"""
    if not seconds:
        return "Unknown Duration"
    
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    else:
        return f"{minutes:02d}:{seconds:02d}"

if __name__ == '__main__':
    # Disable reloader to avoid connection resets while streaming captions
    app.run(debug=True, host='0.0.0.0', port=5000, use_reloader=False)
