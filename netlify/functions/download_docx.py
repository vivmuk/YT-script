import json
import io
import base64

def response_bytes(data: bytes, filename: str, mimetype: str):
    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": mimetype,
            "Content-Disposition": f"attachment; filename={filename}"
        },
        "isBase64Encoded": True,
        "body": base64.b64encode(data).decode('ascii')
    }

def split_into_paragraphs(text: str):
    import re
    parts = re.split(r"\n\n+|(?<=[.!?])\s+(?=[A-Z])", text or '')
    return [p.strip() for p in parts if p.strip()]

def handler(event, context):
    try:
        from docx import Document
        from docx.shared import Pt
        data = json.loads(event.get('body') or '{}')
        transcription = data.get('transcription','')
        title = data.get('title','YouTube Captions')
        duration = data.get('duration','')
        uploader = data.get('uploader','')
        language = data.get('language','')
        caption_type = data.get('caption_type','')

        doc = Document()
        h = doc.add_heading(title, level=0)
        h.style.font.size = Pt(18)
        meta = doc.add_paragraph()
        meta.add_run(f"Uploader: {uploader}\n").bold = True
        meta.add_run(f"Duration: {duration}\n").bold = True
        meta.add_run(f"Language: {language} | Type: {caption_type}\n").bold = True
        doc.add_paragraph("")
        doc.add_heading('Captions', level=1)
        for p in split_into_paragraphs(transcription):
            doc.add_paragraph(p)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe = ''.join(c if c.isalnum() or c=='_' else '_' for c in (title or 'youtube_captions').lower())
        return response_bytes(buf.getvalue(), f"{safe}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return {"statusCode": 500, "headers": {"Content-Type":"application/json"}, "body": json.dumps({"error":"Failed to generate DOCX"})}


